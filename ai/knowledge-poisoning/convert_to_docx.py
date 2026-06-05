#!/usr/bin/env python3
"""将毒知识库载荷从 .md 转换为 .docx 格式"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path("/mnt/d/docxnotes/ai/knowledge-poisoning")
PAYLOAD_DIR = BASE / "payloads_docx"

# 载荷映射: (源md, 目标docx文件名, 标题)
PAYLOADS = [
    # stealth
    ("stealth/security-policy-v4.md",
     "企业安全策略文档_v4.2.docx",
     "企业安全策略文档 v4.2"),

    ("stealth/ai-integration-guide-v3.md",
     "AI系统管理员指南_v3.docx",
     "AI 系统管理员指南 v3"),

    # aggressive
    ("aggressive/emergency-security-update.md",
     "紧急安全更新_2026-06-01.docx",
     "紧急安全更新 — 2026-06-01"),

    ("aggressive/api-config-v2.md",
     "AI服务端API配置_v2.4.docx",
     "AI 服务端 API 配置 v2.4"),

    # research
    ("research/reasoning-framework.md",
     "思维链推理最佳实践_v2.1.docx",
     "思维链推理最佳实践 v2.1"),

    ("research/safety-alignment-experiment.md",
     "AI安全对齐评估_实验记录.docx",
     "AI 安全对齐评估 — 实验记录"),
]

def md_to_docx(md_path: Path, docx_path: Path, title: str):
    """Convert markdown payload to formatted .docx"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 读取 markdown
    md_text = md_path.read_text(encoding="utf-8")

    # 添加标题
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # 逐行处理
    lines = md_text.split("\n")
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()

        # 跳过 frontmatter (---)
        if stripped == "---" and not in_code_block:
            continue

        # 代码块开关
        if stripped.startswith("```"):
            if in_code_block:
                # 结束代码块
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                # 灰色底纹
                from docx.oxml.ns import qn
                shading = p.paragraph_format.element.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F0F0',
                    qn('w:val'): 'clear'
                })
                p.paragraph_format.element.insert(0, shading)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                # 忽略 ``` 后的语言标记
                lang = stripped[3:].strip()
                if lang:
                    p = doc.add_paragraph(f"[{lang}]")
                    p.runs[0].font.size = Pt(8)
                    p.runs[0].font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 空行
        if not stripped:
            doc.add_paragraph("")
            continue

        # 注释行（>）
        if stripped.startswith(">") and not in_code_block:
            p = doc.add_paragraph(stripped[1:].strip())
            p.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            p.runs[0].font.italic = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        # 列表项
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            continue

        if re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s*', '', stripped)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            continue

        # 标题
        if stripped.startswith("##### "):
            doc.add_heading(stripped[6:], level=5)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            # 跳过主标题（已经在 doc 标题中了）
            continue
        else:
            # 普通段落
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    # 保存
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"  ✅ {docx_path.name}  ({docx_path.stat().st_size:,} bytes)")


def main():
    print("🔄 转换载荷到 .docx 格式...\n")

    for rel_path, docx_name, title in PAYLOADS:
        md_file = BASE / "payloads" / rel_path
        docx_file = PAYLOAD_DIR / docx_name

        if not md_file.exists():
            print(f"  ❌ 源文件不存在: {md_file}")
            continue

        md_to_docx(md_file, docx_file, title)

    print(f"\n📁 全部 .docx 文件在: {PAYLOAD_DIR}")
    print(f"📦 共 {len(PAYLOADS)} 个毒文档就绪")

if __name__ == "__main__":
    main()
